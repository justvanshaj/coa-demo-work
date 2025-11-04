import streamlit as st
from docx import Document
import os
import mammoth
import io
from datetime import datetime
import calendar
import random
import pandas as pd

# ---------------------------
# Helper: placeholder replacement preserving style
# ---------------------------
def advanced_replace_text_preserving_style(doc, replacements):
    def replace_in_paragraph(paragraph):
        runs = paragraph.runs
        full_text = ''.join(run.text for run in runs)
        for key, value in replacements.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in full_text:
                new_runs = []
                accumulated = ""
                for run in runs:
                    accumulated += run.text
                    new_runs.append(run)
                    if placeholder in accumulated:
                        style_run = next((r for r in new_runs if placeholder in r.text), new_runs[0])
                        font = style_run.font
                        accumulated = accumulated.replace(placeholder, value)
                        for r in new_runs:
                            r.text = ''
                        if new_runs:
                            new_run = new_runs[0]
                            new_run.text = accumulated
                            try:
                                new_run.font.name = font.name
                            except:
                                pass
                            try:
                                new_run.font.size = font.size
                            except:
                                pass
                            try:
                                new_run.font.bold = font.bold
                            except:
                                pass
                            try:
                                new_run.font.italic = font.italic
                            except:
                                pass
                            try:
                                new_run.font.underline = font.underline
                            except:
                                pass
                            try:
                                new_run.font.color.rgb = font.color.rgb
                            except:
                                pass
                        break

    for para in doc.paragraphs:
        replace_in_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)

# ---------------------------
# DOCX generate / preview
# ---------------------------
def generate_docx(data, template_path="template.docx", output_path="generated_coa.docx"):
    doc = Document(template_path)
    advanced_replace_text_preserving_style(doc, data)
    doc.save(output_path)
    return output_path

def docx_to_html(docx_path):
    with open(docx_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        return result.value

# ---------------------------
# Distribution helper (water-filling)
# ---------------------------
def distribute_within_bounds(target, names, mins, maxs, weights):
    eps = 1e-9
    min_sum = sum(mins[n] for n in names)
    max_sum = sum(maxs[n] for n in names)
    if target + eps < min_sum or target - eps > max_sum:
        raise ValueError("Target not achievable with given bounds.")

    w_sum = sum(weights[n] for n in names)
    if w_sum <= 0:
        vals = {n: target / len(names) for n in names}
    else:
        vals = {n: target * (weights[n] / w_sum) for n in names}

    locked = {n: False for n in names}
    for _ in range(100):
        changed = False
        for n in names:
            if not locked[n]:
                if vals[n] < mins[n]:
                    vals[n] = mins[n]
                    locked[n] = True
                    changed = True
                elif vals[n] > maxs[n]:
                    vals[n] = maxs[n]
                    locked[n] = True
                    changed = True
        if not changed:
            unlocked = [n for n in names if not locked[n]]
            if not unlocked:
                break
            rem = target - sum(vals.values())
            if abs(rem) < 1e-8:
                break
            w_un_sum = sum(weights[n] for n in unlocked)
            if w_un_sum <= 0:
                for n in unlocked:
                    vals[n] += rem / len(unlocked)
            else:
                for n in unlocked:
                    vals[n] += rem * (weights[n] / w_un_sum)

    for n in names:
        vals[n] = max(min(vals[n], maxs[n]), mins[n])
        vals[n] = round(vals[n], 2)

    total = round(sum(vals[n] for n in names), 2)
    diff = round(target - total, 2)
    if abs(diff) >= 0.01:
        adjustable = [n for n in names if mins[n] + 1e-9 < vals[n] < maxs[n] - 1e-9]
        for n in adjustable:
            allow_low = round(vals[n] - mins[n], 2)
            allow_high = round(maxs[n] - vals[n], 2)
            move = max(-allow_low, min(allow_high, diff))
            if abs(move) >= 0.01:
                vals[n] = round(vals[n] + move, 2)
                diff = round(target - sum(vals.values()), 2)
                if abs(diff) < 0.01:
                    break

    final_total = round(sum(vals[n] for n in names), 2)
    if abs(final_total - target) > 0.01:
        raise ValueError("Could not distribute to meet target.")
    return vals

# ---------------------------
# Ranges and mids
# ---------------------------
RANGES = {
    "fat": (0.45, 0.55),
    "air": (2.90, 3.10),
    "ash": (0.45, 0.55),
    "protein": (2.45, 2.55),
    "gum": (80.10, 89.95)
}
MIDS = {k: round((v[0] + v[1]) / 2.0, 4) for k, v in RANGES.items()}

# ---------------------------
# Deterministic & Random calculation
# ---------------------------
def calculate_components_deterministic(moisture):
    remaining = round(100.0 - float(moisture), 4)
    others = ["fat", "air", "ash", "protein"]
    others_mid_sum = sum(MIDS[o] for o in others)
    gum_needed = round(remaining - others_mid_sum, 4)
    gum_min, gum_max = RANGES["gum"]
    if gum_min - 1e-9 <= gum_needed <= gum_max + 1e-9:
        fat = round(MIDS["fat"], 2)
        air = round(MIDS["air"], 2)
        ash = round(MIDS["ash"], 2)
        protein = round(MIDS["protein"], 2)
        gum = round(gum_needed, 2)
        total = round(moisture + fat + air + ash + protein + gum, 2)
        if abs(total - 100.0) > 0.01:
            gum = round(gum + (100.0 - total), 2)
        return gum, protein, ash, air, fat

    gum_try_list = [RANGES["gum"][0], RANGES["gum"][1]]
    for gum_try in gum_try_list:
        available = round(remaining - gum_try, 4)
        mins = {o: RANGES[o][0] for o in others}
        maxs = {o: RANGES[o][1] for o in others}
        weights = {o: MIDS[o] for o in others}
        try:
            allocated = distribute_within_bounds(available, others, mins, maxs, weights)
            fat = allocated["fat"]
            air = allocated["air"]
            ash = allocated["ash"]
            protein = allocated["protein"]
            gum = round(gum_try, 2)
            total = round(moisture + fat + air + ash + protein + gum, 2)
            if abs(total - 100.0) > 0.01:
                residual = round(100.0 - total, 2)
                new_gum = round(gum + residual, 2)
                if RANGES["gum"][0] <= new_gum <= RANGES["gum"][1]:
                    gum = new_gum
            return gum, protein, ash, air, fat
        except ValueError:
            continue
    raise ValueError("Unable to compute deterministic components for this moisture.")

def calculate_components_random(moisture, max_attempts=2000):
    remaining = round(100.0 - float(moisture), 4)
    others = ["fat", "air", "ash", "protein"]
    gum_min, gum_max = RANGES["gum"]

    for attempt in range(max_attempts):
        sampled = {o: round(random.uniform(RANGES[o][0], RANGES[o][1]), 4) for o in others}
        gum_needed = round(remaining - sum(sampled.values()), 4)
        if gum_min - 1e-9 <= gum_needed <= gum_max + 1e-9:
            fat = round(sampled["fat"], 2)
            air = round(sampled["air"], 2)
            ash = round(sampled["ash"], 2)
            protein = round(sampled["protein"], 2)
            gum = round(gum_needed, 2)
            total = round(moisture + fat + air + ash + protein + gum, 2)
            if abs(total - 100.0) > 0.01:
                gum = round(gum + (100.0 - total), 2)
            return gum, protein, ash, air, fat

        gum_try = round(random.uniform(gum_min, gum_max), 4)
        available_for_others = round(remaining - gum_try, 4)
        mins = {o: RANGES[o][0] for o in others}
        maxs = {o: RANGES[o][1] for o in others}
        rand_weights = {o: random.random() + MIDS[o] for o in others}
        try:
            allocated = distribute_within_bounds(available_for_others, others, mins, maxs, rand_weights)
            fat = allocated["fat"]
            air = allocated["air"]
            ash = allocated["ash"]
            protein = allocated["protein"]
            gum = round(gum_try, 2)
            total = round(moisture + fat + air + ash + protein + gum, 2)
            if abs(total - 100.0) > 0.01:
                residual = round(100.0 - total, 2)
                new_gum = round(gum + residual, 2)
                if gum_min <= new_gum <= gum_max:
                    gum = new_gum
            return gum, protein, ash, air, fat
        except ValueError:
            continue

    raise ValueError("Randomized sampling failed to find feasible components; try different moisture or increase attempts.")

# ---------------------------
# Streamlit UI
# ---------------------------
st.set_page_config(page_title="COA Generator", layout="wide")
st.title("🧪 COA Document Generator (Auto Randomized Components)")

# session state keys
if "components" not in st.session_state:
    st.session_state["components"] = None
if "components_moisture" not in st.session_state:
    st.session_state["components_moisture"] = None

# Input area (use normal widgets so buttons outside can control behavior)
col_left, col_right = st.columns([2, 1])
with col_left:
    code = st.selectbox(
        "Select Product Code Range",
        [f"{i}-{i+500}" for i in range(500, 10001, 500)]
    )
    st.info(f"📄 Using template: COA {code}.docx")

    date = st.text_input("Date (e.g., July 2025)")

    # Auto Best Before
    best_before = ""
    if date:
        parsed = None
        for fmt_try in ("%B %Y", "%b %Y", "%B, %Y", "%b, %Y"):
            try:
                parsed = datetime.strptime(date.strip().title(), fmt_try)
                break
            except:
                try:
                    parsed = datetime.strptime(date.strip(), fmt_try)
                    break
                except:
                    parsed = None
        if parsed:
            year = parsed.year + 2
            month = parsed.month - 1
            if month == 0:
                month = 12
                year -= 1
            best_before = f"{calendar.month_name[month].upper()} {year}"
            st.success(f"Best Before auto-filled: {best_before}")
        else:
            st.warning("Enter Date in format: July 2025")

    batch_no = st.text_input("Batch Number")
    moisture = st.number_input("Moisture (%)", min_value=0.0, max_value=99.0, step=0.01, value=10.00, format="%.2f")
    st.markdown(
        "Component values (Fat, Air, Ash, Protein, Gum) are **automatically randomized** (within specified ranges) "
        "and will sum with Moisture to 100%."
    )

with col_right:
    ph = st.text_input("pH Level (e.g., 6.7)")
    mesh_200 = st.text_input("200 Mesh (%)")
    viscosity_2h = st.text_input("Viscosity After 2 Hours (CPS)")
    viscosity_24h = st.text_input("Viscosity After 24 Hours (CPS)")

# regenerate automatically when moisture changes (to avoid stale components)
def ensure_components_for_current_moisture():
    if st.session_state["components"] is None or st.session_state["components_moisture"] != round(moisture, 2):
        try:
            gum, protein, ash, air, fat = calculate_components_random(moisture)
        except Exception:
            # fallback deterministic if random fails
            gum, protein, ash, air, fat = calculate_components_deterministic(moisture)
        st.session_state["components"] = {
            "Moisture": round(moisture, 2),
            "Fat": float(f"{fat:.2f}"),
            "Air": float(f"{air:.2f}"),
            "Ash": float(f"{ash:.2f}"),
            "Protein": float(f"{protein:.2f}"),
            "Gum": float(f"{gum:.2f}")
        }
        st.session_state["components_moisture"] = round(moisture, 2)

ensure_components_for_current_moisture()

# Buttons row: Refresh components and Generate COA side-by-side
btn_col1, btn_col2 = st.columns([1, 1])
with btn_col1:
    refresh = st.button("🔄 Refresh components")
with btn_col2:
    generate = st.button("Generate COA")

# If refresh clicked, randomize and update session_state
if refresh:
    try:
        gum, protein, ash, air, fat = calculate_components_random(moisture)
        st.session_state["components"] = {
            "Moisture": round(moisture, 2),
            "Fat": float(f"{fat:.2f}"),
            "Air": float(f"{air:.2f}"),
            "Ash": float(f"{ash:.2f}"),
            "Protein": float(f"{protein:.2f}"),
            "Gum": float(f"{gum:.2f}")
        }
        st.session_state["components_moisture"] = round(moisture, 2)
        st.success("Components refreshed (randomized).")
    except Exception as e:
        st.error(f"Random generation failed: {e}")

# Show composition summary as a single-row table (columns horizontally)
if st.session_state["components"]:
    comp = st.session_state["components"]
    summary_df = pd.DataFrame([{
        "Moisture (%)": f"{comp['Moisture']:.2f}",
        "Fat (%)": f"{comp['Fat']:.2f}",
        "Air (%)": f"{comp['Air']:.2f}",
        "Ash (%)": f"{comp['Ash']:.2f}",
        "Protein (%)": f"{comp['Protein']:.2f}",
        "Gum (%)": f"{comp['Gum']:.2f}"
    }])
    st.subheader("Composition summary (current)")
    st.dataframe(summary_df, width=900)
    total = round(float(comp["Moisture"]) + float(comp["Fat"]) + float(comp["Air"]) + float(comp["Ash"]) + float(comp["Protein"]) + float(comp["Gum"]), 2)
    st.info(f"Total = {total:.2f}%")

# Handle Generate COA click: validate and generate
if generate:
    if st.session_state["components"] is None:
        st.error("Components not available. Try refreshing.")
    else:
        comp = st.session_state["components"]
        total = round(float(comp["Moisture"]) + float(comp["Fat"]) + float(comp["Air"]) + float(comp["Ash"]) + float(comp["Protein"]) + float(comp["Gum"]), 2)
        if abs(total - 100.0) > 0.01:
            st.error(f"Current components total {total:.2f}% does not equal 100.00%. Refresh components.")
        else:
            data = {
                "DATE": date,
                "BATCH_NO": batch_no,
                "BEST_BEFORE": best_before,
                "MOISTURE": f"{comp['Moisture']:.2f}%",
                "PH": ph,
                "MESH_200": f"{mesh_200}%",
                "VISCOSITY_2H": viscosity_2h,
                "VISCOSITY_24H": viscosity_24h,
                "GUM_CONTENT": f"{comp['Gum']:.2f}%",
                "PROTEIN": f"{comp['Protein']:.2f}%",
                "ASH_CONTENT": f"{comp['Ash']:.2f}%",
                "AIR": f"{comp['Air']:.2f}%",
                "FAT": f"{comp['Fat']:.2f}%"
            }

            template_path = f"COA {code}.docx"
            output_path = "generated_coa.docx"

            if not os.path.exists(template_path):
                st.error(f"Template file 'COA {code}.docx' not found.")
            else:
                generate_docx(data, template_path=template_path, output_path=output_path)

                try:
                    html = docx_to_html(output_path)
                    st.subheader("📄 Preview")
                    st.components.v1.html(f"<div style='padding:15px'>{html}</div>", height=700, scrolling=True)
                except Exception:
                    st.warning("Preview failed. You can still download the file below.")

                safe_batch = (batch_no or "batch").replace("/", "_").replace("\\", "_").replace(" ", "_")
                final_filename = f"COA-{safe_batch}-{code}.docx"

                with open(output_path, "rb") as f:
                    doc_bytes = f.read()
                buffer = io.BytesIO(doc_bytes)

                st.download_button(
                    label="📥 Download COA (DOCX)",
                    data=buffer,
                    file_name=final_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
